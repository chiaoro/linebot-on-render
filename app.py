from flask import Flask, request, abort, jsonify
from linebot import LineBotApi, WebhookHandler
from linebot.exceptions import InvalidSignatureError, LineBotApiError
from linebot.models import TextMessage, MessageEvent, TextSendMessage, FlexSendMessage
import os, json, tempfile, requests, mimetypes, smtplib, gspread
from email.mime.text import MIMEText
from oauth2client.service_account import ServiceAccountCredentials
from datetime import datetime
from dotenv import load_dotenv
from utils.line_push import push_text_to_user
from utils.schedule_utils import handle_submission
from utils.google_auth import get_gspread_client
from utils.google_sheets import log_meeting_reply, get_doctor_name
from utils.state_manager import set_state, get_state, clear_state
import re
from meeting_reminder import send_meeting_reminder
from monthly_reminder import send_monthly_fixed_reminders
from event_reminder import send_important_event_reminder
from daily_notifier import run_daily_push
from utils.night_shift_fee import handle_night_shift_request
from utils.night_shift_fee_generator import run_generate_night_fee_word
from utils.night_shift_fee import daily_night_fee_reminder
from meeting_leave import handle_meeting_leave_response
from meeting_leave_scheduler import run_meeting_leave_scheduler

# ✅ 初始化 Flask 和 LineBot API
app = Flask(__name__)
line_bot_api = LineBotApi(os.getenv("LINE_CHANNEL_ACCESS_TOKEN"))
handler = WebhookHandler(os.getenv("LINE_CHANNEL_SECRET"))

# ✅ Google Sheets 設置
SHEET_URL = "https://docs.google.com/spreadsheets/d/1fHf5XlbvLMd6ytAh_t8Bsi5ghToiQHZy1NlVfEG7VIo/edit"
WORKSHEET_NAME = "夜點費申請"

# ✅ Google Sheets 認證
SCOPE = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
CREDS_DICT = json.loads(os.getenv("GOOGLE_CREDENTIALS"))
CREDS = ServiceAccountCredentials.from_json_keyfile_dict(CREDS_DICT, SCOPE)
GC = gspread.authorize(CREDS)

# ✅ 主選單 Flex 定義
def get_main_menu():
    return FlexSendMessage("主選單", {
        "type": "bubble",
        "body": {
            "type": "box", "layout": "vertical", "contents": [
                {"type": "text", "text": "📋 請選擇服務類別", "weight": "bold", "size": "lg", "margin": "md"},
                *[
                    {"type": "button", "action": {"type": "message", "label": label, "text": label},
                     "style": "primary", "margin": "md"}
                    for label in ["門診調整服務", "值班調整服務", "支援醫師服務", "新進醫師服務", "其他表單服務"]
                ]
            ]
        }
    })

# ✅ 處理夜點費申請
@app.route("/generate-night-fee-word", methods=["GET"])
def generate_night_fee_word():
    try:
        from utils.night_shift_fee_generator import run_generate_night_fee_word
        run_generate_night_fee_word()
        return "✅ 夜點費申請表產生完成", 200
    except Exception as e:
        return f"❌ 夜點費申請表產生錯誤：{e}", 500


# ✅ 處理夜點費申請
def handle_night_shift_request(user_id, user_msg):
    sheet = GC.open_by_url(SHEET_URL).worksheet(WORKSHEET_NAME)
    now = datetime.now().strftime("%Y/%m/%d %H:%M:%S")
    doctor_name = user_msg.replace("夜點費", "").strip()
    sheet.append_row([now, doctor_name, "未提醒"])
    return f"已收到 {doctor_name} 的夜點費申請，將於每月 1~5 號進行催繳提醒。"

# ✅ 回應 LINE 訊息
@handler.add(MessageEvent, message=TextMessage)
def handle_message(event):
    user_id = event.source.user_id
    user_msg = event.message.text.strip()

    # ✅ 夜點費申請
    if "夜點費" in user_msg:
        reply = handle_night_shift_request(user_id, user_msg)
        if reply:
            line_bot_api.reply_message(event.reply_token, TextSendMessage(text=reply))
        return

    # ✅ 主選單顯示
    if user_msg == "主選單":
        line_bot_api.reply_message(event.reply_token, get_main_menu())
        return

    # ✅ 其他訊息處理
    line_bot_api.reply_message(event.reply_token, TextSendMessage(text="⚠️ 無效指令，請輸入「主選單」重新開始。"))

# ✅ LINE Webhook
@app.route("/callback", methods=["POST"])
def callback():
    signature = request.headers.get('X-Line-Signature')
    body = request.get_data(as_text=True)
    try:
        handler.handle(body, signature)
    except InvalidSignatureError:
        abort(400)
    return "OK"



# ✅ 夜點費申請生成 Word 文件
@app.route("/generate-night-fee-word", methods=["GET"])
def generate_night_fee_word():
    try:
        from utils.night_shift_fee_generator import run_generate_night_fee_word
        run_generate_night_fee_word()
        return "✅ 夜點費申請表產生完成", 200
    except Exception as e:
        return f"❌ 夜點費申請表產生錯誤：{e}", 500

# ✅ 生成夜點費申請表的函式
def run_generate_night_fee_word():
    sheet = GC.open_by_url(SHEET_URL).worksheet(WORKSHEET_NAME)
    records = sheet.get_all_records()

    # 建立 Word 文件
    from docx import Document
    doc = Document()
    doc.add_heading('夜點費申請表', 0)

    for record in records:
        doctor_name = record['醫師姓名']
        doc.add_paragraph(f"醫師：{doctor_name}")

    # 存儲為 Word 文件
    file_path = "/mnt/data/night_fee_request.docx"
    doc.save(file_path)
    return file_path



# ✅ 會議請假處理
@app.route("/meeting-leave", methods=["POST"])
def meeting_leave():
    data = request.get_json()
    user_id = data.get("user_id")
    leave_reason = data.get("leave_reason")

    if not user_id or not leave_reason:
        return jsonify({"status": "error", "message": "缺少欄位"}), 400

    doctor_name = get_doctor_name(DOCTOR_SHEET_URL, user_id)
    log_meeting_reply(RECORD_SHEET_URL, user_id, doctor_name, "請假", leave_reason)

    return jsonify({"status": "success", "message": f"{doctor_name} 的請假已成功記錄。"}), 200

# ✅ 每月夜點費提醒
@app.route("/night-shift-reminder", methods=["GET"])
def night_shift_reminder():
    try:
        daily_night_fee_reminder()
        return "✅ 夜點費每日提醒完成", 200
    except Exception as e:
        return f"❌ 夜點費提醒錯誤：{e}", 500

# ✅ 每月固定提醒推播
@app.route("/monthly-reminder", methods=["GET"])
def monthly_reminder():
    try:
        send_monthly_fixed_reminders()
        return "✅ 固定日期推播完成", 200
    except Exception as e:
        return f"❌ 固定日期推播錯誤：{e}", 500

# ✅ 重要會議提醒推播
@app.route("/event-reminder", methods=["GET"])
def event_reminder():
    try:
        send_important_event_reminder()
        return "✅ 重要會議提醒完成", 200
    except Exception as e:
        return f"❌ 重要會議提醒錯誤：{e}", 500

# ✅ 每日推播
@app.route("/daily-push", methods=["GET"])
def daily_push():
    try:
        run_daily_push()
        return "✅ 今日推播完成", 200
    except Exception as e:
        return f"❌ 今日推播錯誤：{e}", 500

# ✅ 喚醒 Bot
@app.route("/ping", methods=["GET"])
def ping():
    return "Bot is awake!", 200




# ✅ 生成夜點費申請表的函式
def run_generate_night_fee_word():
    # 取得夜點費申請表的 Google Sheets 資料
    sheet = GC.open_by_url(SHEET_URL).worksheet(WORKSHEET_NAME)
    records = sheet.get_all_records()

    # 建立 Word 文件
    from docx import Document
    doc = Document()
    doc.add_heading('夜點費申請表', 0)

    # 根據醫師資料生成申請表內容
    for record in records:
        doctor_name = record['醫師姓名']
        doc.add_paragraph(f"醫師：{doctor_name}")

    # 儲存 Word 檔案
    file_path = "/mnt/data/night_fee_request.docx"
    doc.save(file_path)

    return file_path

# ✅ 每月夜點費提醒（新增根據科別生成 Word 範本）
def daily_night_fee_reminder():
    today = date.today()
    if not (1 <= today.day <= 5):  # 只有在每月1到5號進行提醒
        return

    sheet = GC.open_by_url(SHEET_URL).worksheet(WORKSHEET_NAME)
    records = sheet.get_all_records()

    for idx, rec in enumerate(records, start=2):
        apply_time = rec.get("時間", "")
        doctor = rec.get("醫師姓名")
        status = rec.get("提醒狀態")
        # 檢查是否為上個月且未提醒
        try:
            apply_date = datetime.strptime(apply_time, "%Y/%m/%d %H:%M:%S").date()
        except:
            continue
        last_month = today.month - 1 or 12
        if apply_date.month == last_month and status != "已提醒":
            text = f"📌 {doctor}，請於本月 1~5 號繳交 {apply_date.strftime('%Y/%m')} 夜點費資料，謝謝！"
            push_text_to_group(GROUP_ID, text)
            sheet.update_cell(idx, list(records[0].keys()).index("提醒狀態")+1, "已提醒")


# ✅ 處理夜點費申請的回應
@app.route("/submit-night-fee", methods=["POST"])
def submit_night_fee():
    data = request.get_json()
    doctor_name = data.get("doctor_name")
    if not doctor_name:
        return jsonify({"status": "error", "message": "缺少醫師姓名"}), 400

    # 將資料寫入夜點費申請表
    sheet = GC.open_by_url(SHEET_URL).worksheet(WORKSHEET_NAME)
    now = datetime.now().strftime("%Y/%m/%d %H:%M:%S")
    sheet.append_row([now, doctor_name, "未提醒"])

    # 發送確認回覆
    push_text_to_user(event.reply_token, f"已收到 {doctor_name} 的夜點費申請，將於每月 1~5 號進行催繳提醒。")
    return jsonify({"status": "success", "message": f"{doctor_name} 的夜點費申請已成功提交。"}), 200


# ✅ 錯誤處理
@app.route("/error-handler", methods=["GET"])
def error_handler():
    try:
        # 測試用錯誤，根據需求處理其他功能
        raise Exception("測試錯誤")
    except Exception as e:
        return f"❌ 錯誤發生：{str(e)}", 500


# ✅ 主選單處理邏輯
@app.route("/main-menu", methods=["GET"])
def main_menu():
    try:
        # 主選單回應
        line_bot_api.reply_message(
            event.reply_token,
            get_main_menu()
        )
        return "✅ 主選單已顯示", 200
    except Exception as e:
        return f"❌ 顯示主選單時發生錯誤：{str(e)}", 500


# ✅ 重啟 Bot
@app.route("/restart-bot", methods=["POST"])
def restart_bot():
    try:
        # 重啟流程（可根據需求修改）
        return jsonify({"status": "success", "message": "Bot 已重新啟動"}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ✅ 簡單錯誤回應範例（Debug）
@app.route("/debug", methods=["GET"])
def debug():
    try:
        # 模擬錯誤情境
        result = 10 / 0  # Division by zero for error simulation
        return jsonify({"result": result}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500



# ✅ 值班調整處理
@app.route("/duty-swap", methods=["POST"])
def duty_swap():
    data = request.get_json()
    original_doctor = data.get("original_doctor")
    swap_doctor = data.get("swap_doctor")
    swap_date = data.get("swap_date")
    reason = data.get("reason")
    
    if not original_doctor or not swap_doctor or not swap_date or not reason:
        return jsonify({"status": "error", "message": "缺少必填欄位"}), 400
    
    # 寫入 Google Sheets
    duty_sheet = GC.open_by_url(SHEET_URL).worksheet("值班調整")
    now = datetime.now().strftime("%Y/%m/%d %H:%M:%S")
    duty_sheet.append_row([now, original_doctor, swap_doctor, swap_date, reason, "未處理"])

    # 發送確認回覆
    push_text_to_user(event.reply_token, f"已收到您的值班調整申請，原醫師：{original_doctor}，調換醫師：{swap_doctor}，日期：{swap_date}，原因：{reason}。")
    return jsonify({"status": "success", "message": "值班調整申請已成功提交。"}), 200

# ✅ 值班代理處理
@app.route("/duty-proxy", methods=["POST"])
def duty_proxy():
    data = request.get_json()
    original_doctor = data.get("original_doctor")
    proxy_doctor = data.get("proxy_doctor")
    proxy_date = data.get("proxy_date")
    reason = data.get("reason")
    
    if not original_doctor or not proxy_doctor or not proxy_date or not reason:
        return jsonify({"status": "error", "message": "缺少必填欄位"}), 400
    
    # 寫入 Google Sheets
    duty_sheet = GC.open_by_url(SHEET_URL).worksheet("值班代理")
    now = datetime.now().strftime("%Y/%m/%d %H:%M:%S")
    duty_sheet.append_row([now, original_doctor, proxy_doctor, proxy_date, reason, "未處理"])

    # 發送確認回覆
    push_text_to_user(event.reply_token, f"已收到您的值班代理申請，原醫師：{original_doctor}，代理醫師：{proxy_doctor}，日期：{proxy_date}，原因：{reason}。")
    return jsonify({"status": "success", "message": "值班代理申請已成功提交。"}), 200



# ✅ 值班調整通知
@app.route("/duty-swap-reminder", methods=["GET"])
def duty_swap_reminder():
    # 讀取值班調整資料
    duty_sheet = GC.open_by_url(SHEET_URL).worksheet("值班調整")
    records = duty_sheet.get_all_records()
    
    for record in records:
        if record.get("狀態") != "已處理":
            text = f"📌 值班調整提醒：\n原醫師：{record['原醫師']}\n調換醫師：{record['調換醫師']}\n日期：{record['日期']}\n原因：{record['原因']}"
            push_text_to_group(GROUP_ID, text)
            # 更新狀態為已處理
            duty_sheet.update_cell(records.index(record) + 2, list(records[0].keys()).index("狀態") + 1, "已處理")
    return "✅ 值班調整提醒完成", 200

# ✅ 每月夜點費提醒
@app.route("/night-shift-fee-reminder", methods=["GET"])
def night_shift_fee_reminder():
    sheet = GC.open_by_url(SHEET_URL).worksheet("夜點費")
    records = sheet.get_all_records()

    for record in records:
        doctor_name = record.get("醫師姓名")
        status = record.get("提醒狀態")
        if status != "已提醒":
            text = f"📌 {doctor_name}，請於本月繳交夜點費。"
            push_text_to_group(GROUP_ID, text)
            # 更新提醒狀態為已提醒
            sheet.update_cell(records.index(record) + 2, list(records[0].keys()).index("提醒狀態") + 1, "已提醒")
    return "✅ 夜點費提醒完成", 200



from docx import Document

# ✅ 根據醫師科別生成夜點費申請表的 Word 檔案
def generate_night_fee_word_by_specialty():
    sheet = GC.open_by_url(SHEET_URL).worksheet("夜點費")
    records = sheet.get_all_records()

    # 根據科別生成不同格式的 Word 文件
    for record in records:
        doctor_name = record.get("醫師姓名")
        specialty = record.get("科別")
        
        # 根據科別選擇對應的 Word 模板
        doc = Document()
        doc.add_heading(f"{specialty} - 夜點費申請表", 0)

        doc.add_paragraph(f"醫師姓名: {doctor_name}")
        doc.add_paragraph(f"科別: {specialty}")
        doc.add_paragraph(f"申請時間: {datetime.now().strftime('%Y/%m/%d %H:%M:%S')}")

        # 保存文件
        file_path = f"/mnt/data/{doctor_name}_night_fee_request.docx"
        doc.save(file_path)

        # 上傳到 Google Drive 或其他地方進行存儲或後續處理
        # 假設有上傳函式
        upload_to_drive(file_path)
        
    return "✅ 夜點費申請表生成完成", 200


# ✅ 上傳檔案至 Google Drive 的函式
def upload_to_drive(file_path):
    # 假設使用 Google Drive API 上傳文件
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    # 構建 Google Drive 服務
    drive_service = build('drive', 'v3', credentials=CREDS)

    # 設定檔案名稱與上傳的目錄
    file_metadata = {'name': os.path.basename(file_path), 'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'}
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # 上傳文件
    file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()

    return file.get('id')




# ✅ 錯誤處理
@app.route("/error-handler", methods=["GET"])
def error_handler():
    try:
        # 測試用錯誤，根據需求處理其他功能
        raise Exception("測試錯誤")
    except Exception as e:
        return f"❌ 錯誤發生：{str(e)}", 500


# ✅ 每月夜點費提醒
@app.route("/night-shift-fee-reminder", methods=["GET"])
def night_shift_fee_reminder():
    sheet = GC.open_by_url(SHEET_URL).worksheet("夜點費")
    records = sheet.get_all_records()

    for record in records:
        doctor_name = record.get("醫師姓名")
        status = record.get("提醒狀態")
        if status != "已提醒":
            text = f"📌 {doctor_name}，請於本月繳交夜點費。"
            push_text_to_group(GROUP_ID, text)
            # 更新提醒狀態為已提醒
            sheet.update_cell(records.index(record) + 2, list(records[0].keys()).index("提醒狀態") + 1, "已提醒")
    return "✅ 夜點費提醒完成", 200



# ✅ 定期排程功能 - 夜點費提醒（每月1日提醒）
@app.route("/night-shift-fee-reminder-monthly", methods=["GET"])
def night_shift_fee_reminder_monthly():
    try:
        # 呼叫夜點費提醒功能
        night_shift_fee_reminder()
        return "✅ 夜點費提醒成功發送", 200
    except Exception as e:
        return f"❌ 發送失敗：{str(e)}", 500

# ✅ 排程測試 - 每日自動排程
@app.route("/daily-check-schedule", methods=["GET"])
def daily_check_schedule():
    try:
        # 這裡是處理每天的排程（如每天檢查是否有新資料）
        run_daily_push()  # 假設你有一個每天推播的任務
        return "✅ 每日排程檢查完成", 200
    except Exception as e:
        return f"❌ 排程檢查失敗：{str(e)}", 500

# ✅ 測試路由 - 檢查系統是否正常運行
@app.route("/ping", methods=["GET"])
def ping():
    return "Bot is awake!", 200



# ✅ 申請請假 - 醫師可以申請是否參加院務會議
@app.route("/meeting-leave", methods=["POST"])
def meeting_leave():
    data = request.get_json()
    doctor_name = data.get("doctor_name")
    leave_status = data.get("leave_status")
    
    if not doctor_name or leave_status not in ["出席", "請假"]:
        return jsonify({"status": "error", "message": "請提供完整資料"}), 400
    
    # 醫師出席或請假，將資料寫入 Google Sheets
    sheet = GC.open_by_url(SHEET_URL).worksheet("院務會議請假")
    now = datetime.now().strftime("%Y/%m/%d %H:%M:%S")
    sheet.append_row([now, doctor_name, leave_status])
    
    # 發送確認回覆
    push_text_to_user(event.reply_token, f"已收到您的請假申請，狀態：{leave_status}")
    return jsonify({"status": "success", "message": "請假申請成功"}), 200

# ✅ 值班調整 - 醫師申請值班調換
@app.route("/duty-swap", methods=["POST"])
def duty_swap():
    data = request.get_json()
    original_doctor = data.get("original_doctor")
    swap_doctor = data.get("swap_doctor")
    swap_date = data.get("swap_date")
    reason = data.get("reason")
    
    if not original_doctor or not swap_doctor or not swap_date or not reason:
        return jsonify({"status": "error", "message": "請提供完整資料"}), 400
    
    # 醫師值班調整，將資料寫入 Google Sheets
    sheet = GC.open_by_url(SHEET_URL).worksheet("值班調整")
    now = datetime.now().strftime("%Y/%m/%d %H:%M:%S")
    sheet.append_row([now, original_doctor, swap_doctor, swap_date, reason, "未處理"])
    
    # 發送確認回覆
    push_text_to_user(event.reply_token, f"已收到您的值班調整申請，原醫師：{original_doctor}，調換醫師：{swap_doctor}，日期：{swap_date}，原因：{reason}。")
    return jsonify({"status": "success", "message": "值班調整申請已成功提交"}), 200





# ✅ 啟動 Flask 服務
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))  # 取得或預設使用端口
    print(f"✅ Flask app starting on port {port}")
    app.run(host="0.0.0.0", port=port)
